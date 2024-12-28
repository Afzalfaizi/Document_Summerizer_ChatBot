from fastapi import FastAPI, File, UploadFile, Query
from typing_extensions import TypedDict
from typing import Annotated
from langgraph.graph import StateGraph, START, END
from langgraph.graph.message import add_messages
from langgraph.checkpoint.memory import MemorySaver
from langchain_google_genai import GoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import FAISS  # Updated FAISS import
from langchain.text_splitter import CharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain.schema import HumanMessage, AIMessage
from langchain.docstore.document import Document
from langchain.tools import tool  # Added missing import
from pdfplumber import open as open_pdf
from docx import Document as DocxDocument
from dotenv import load_dotenv
import os

# Load environment variables from a .env file
load_dotenv()

# Initialize LLM and Embedding Models with API keys from environment variables
llm = GoogleGenerativeAI(model="gemini-1.5-flash", google_api_key=os.getenv("GOOGLE_API_KEY"))
embeddings_model = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=os.getenv("GOOGLE_API_KEY"))

# Initialize memory store and checkpointer for state management
memory_store = {}
checkpointer = MemorySaver()

# Define message state using TypedDict for type hinting
class MessagesState(TypedDict):
    messages: Annotated[list, add_messages]

# Tools

def extract_text_from_pdf(file):
    """Extract text from PDF file."""
    # Open the PDF file and extract text from each page
    with open_pdf(file) as pdf:
        text = ''.join([page.extract_text() for page in pdf.pages])
    return text

def extract_text_from_docx(file):
    """Extract text from DOCX file."""
    # Open the DOCX file and extract text from each paragraph
    doc = DocxDocument(file)
    return "\n".join([para.text for para in doc.paragraphs])

@tool
def document_uploader(file_path: str) -> str:
    """Uploads and processes a document with optimized embedding generation."""
    # Determine the file extension to handle different file types
    file_extension = file_path.split('.')[-1].lower()
    try:
        # Extract text based on file type
        if file_extension == "txt":
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
        elif file_extension == "pdf":
            content = extract_text_from_pdf(file_path)
        elif file_extension == "docx":
            content = extract_text_from_docx(file_path)
        else:
            return "Unsupported file format. Please upload a .txt, .pdf, or .docx file."

        # Split text into chunks for embedding generation
        text_splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        chunks = text_splitter.split_text(content)

        # Generate embeddings for the text chunks
        documents = [Document(page_content=chunk) for chunk in chunks]
        embeddings = embeddings_model.embed_documents([doc.page_content for doc in documents])

        # Store embeddings in FAISS vector store
        vector_store = FAISS.from_documents(documents, embeddings_model)

        # Store the document content in memory_store
        memory_store["uploaded_document"] = content

        return "Document uploaded and processed successfully!"
    except Exception as e:
        return f"Error processing file: {str(e)}"

@tool
def document_retriever(query: str, documents: list) -> str:
    """Retrieves the most relevant chunks from documents."""
    # Split documents into chunks for retrieval
    text_splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    chunks = [Document(page_content=chunk) for doc in documents for chunk in text_splitter.split_text(doc)]
    # Generate embeddings for individual chunks
    embeddings = embeddings_model.embed_documents([chunk.page_content for chunk in chunks])
    vector_store = FAISS.from_documents(chunks, embeddings_model)
    retriever = vector_store.as_retriever(search_type="similarity", search_kwargs={"k": 3})
    results = retriever.get_relevant_documents(query)
    return "\n".join([res.page_content for res in results])

@tool
def summarizer(document: str) -> str:
    """Summarizes the document."""
    prompt = f"Summarize the following document: {document[:2000]}"
    summary = llm.invoke([prompt])
    return summary

@tool
def question_answering(query: str, retriever) -> str:
    """Answers user queries based on document content."""
    qa_chain = RetrievalQA.from_chain_type(llm, retriever=retriever)
    answer = qa_chain.run(query)
    return answer

@tool
def memory_tool(thread_id: str, message: str) -> str:
    """Stores and retrieves conversation context."""
    if thread_id not in memory_store:
        memory_store[thread_id] = []
    memory_store[thread_id].append(message)
    return f"Stored message for thread {thread_id}. Current history: {memory_store[thread_id]}"

# Define the assistant node
def assistant(state: MessagesState):
    try:
        user_message = state["messages"][-1]  # Get the last message (HumanMessage object)
        if isinstance(user_message, HumanMessage):
            # Process the user message using the LLM
            response = llm.invoke([user_message.content])  # Access content attribute
            return {"messages": [AIMessage(content=response)]}
        else:
            # Log the invalid message type for debugging
            print(f"Invalid message type: {type(user_message)}. Expected HumanMessage.")
            raise ValueError("Invalid message type for assistant processing.")
    except Exception as e:
        return {"error": f"Error in assistant node: {str(e)}"}

# Define the document summarizer node
def summarize_document(state: MessagesState):
    doc_message = state["messages"][0]  # Get the first message (HumanMessage object)
    if isinstance(doc_message, HumanMessage):
        doc_text = doc_message.content  # Access content attribute
        text_splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        chunks = text_splitter.split_text(doc_text)
        documents = [Document(page_content=chunk) for chunk in chunks]
        embeddings = embeddings_model.embed_documents([doc.page_content for doc in documents])
        vector_store = FAISS.from_documents(documents, embeddings_model)
        retriever = vector_store.as_retriever(search_type="similarity", search_kwargs={"k": 3})
        qa_chain = RetrievalQA.from_chain_type(llm, retriever=retriever)
        summary = qa_chain.run("Summarize the document.")
        return {"messages": [AIMessage(content=summary)]}
    else:
        raise ValueError("Invalid message type for document summarization.")

# Build the graph for state management
builder = StateGraph(MessagesState)
builder.add_node("assistant", assistant)
builder.add_node("summarize", summarize_document)
builder.add_edge(START, "summarize")
builder.add_edge("summarize", "assistant")
builder.add_edge("assistant", END)

# Compile graph with memory checkpointer
graph = builder.compile(checkpointer=checkpointer)

# Create FastAPI app
app = FastAPI()

@app.post("/upload/")
async def upload_document(file: UploadFile = File(...)):
    """
    Endpoint to upload a document, process it, and store its content and summary.
    """
    try:
        # Check file extension
        file_extension = os.path.splitext(file.filename)[-1].lower()
        # Read and process the file based on its type
        if file_extension == ".txt":
            content = await file.read()
            text = content.decode("utf-8")
        elif file_extension == ".pdf":
            text = extract_text_from_pdf(file.file)
        elif file_extension == ".docx":
            text = extract_text_from_docx(file.file)
        else:
            return {"error": "Unsupported file format. Please upload a .txt, .pdf, or .docx file."}
        
        # Store the document content in memory_store
        memory_store["uploaded_document"] = text
        
        # Summarize the document and store the summary
        summary = summarizer(text)
        memory_store["document_summary"] = summary
        
        return {"message": "Document uploaded and summarized successfully!", "summary": summary}
    except Exception as e:
        return {"error": str(e)}

@app.get("/chat/{query}")
def get_content(query: str, thread_id: str = Query(...)):
    """
    Handles user queries with dynamic thread ID for memory context.
    """
    try:
        # Retrieve the document content from memory_store
        if "uploaded_document" not in memory_store:
            return {"error": "No document uploaded. Please upload a document."}
        
        document_content = memory_store["uploaded_document"]
        text_splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        chunks = text_splitter.split_text(document_content)
        documents = [Document(page_content=chunk) for chunk in chunks]
        embeddings = embeddings_model.embed_documents([doc.page_content for doc in documents])
        vector_store = FAISS.from_documents(documents, embeddings_model)
        retriever = vector_store.as_retriever(search_type="similarity", search_kwargs={"k": 3})
        
        # Answer the query using the document content
        qa_chain = RetrievalQA.from_chain_type(llm, retriever=retriever)
        answer = qa_chain.run(query)
        
        # Store the question and answer in memory_store under the given thread_id
        if thread_id not in memory_store:
            memory_store[thread_id] = []
        memory_store[thread_id].append({"question": query, "answer": answer})
        
        return {"answer": answer}
    except Exception as e:
        return {"error": str(e)}

@app.get("/history/{thread_id}")
def get_history(thread_id: str):
    """
    Retrieves the history of questions and answers for a given thread ID.
    """
    try:
        # Check if the thread_id exists in memory_store
        if thread_id not in memory_store:
            return {"error": "No history found for the given thread ID."}
        
        # Retrieve the history for the given thread_id
        history = memory_store[thread_id]
        return {"history": history}
    except Exception as e:
        return {"error": str(e)}